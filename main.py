import os
import subprocess
import pathlib
from PIL import Image, ImageOps
import pygetwindow
import time
import pyautogui
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import Inches
def doc_creator(res):
    doc = docx.Document()
    text_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    for i in res:
        head = doc.add_heading("Assignment-"+ str(i), 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cnt = 1
        for j in res[i]:
            doc.add_paragraph(f"Q{cnt})")
            doc.add_paragraph().add_run("Source Code:- ").bold = True
            with open(j[0], 'r') as f:
                doc.add_paragraph(f.read())
            print(j[1])
            doc.add_paragraph().add_run("Output:- ").bold = True
            doc.add_picture(j[1], width = text_width, height=Inches(3))
            cnt += 1
    doc.save("doc1.docx")

def img(ass_no, idx):
    titles = pygetwindow.getAllTitles()
    for i in titles:
        if(i.endswith("py.exe")):
            window = pygetwindow.getWindowsWithTitle(i)[0]
            break
    left, top = window.topleft
    right, bottom = window.bottomright
    pyautogui.screenshot(f'screenshot{ass_no}({idx}).jpg')
    im = Image.open(f'screenshot{ass_no}({idx}).jpg')
    im = im.crop((left, top, right, bottom))

    im = ImageOps.invert(im)
    im.save(f'screenshot{ass_no}({idx}).jpg')


def run_c_program(ass_no, idx, program_name='program.c'):
    # Compile the C-program
    os.system('cls')
    print(program_name, i)
    subprocess.run(['gcc', program_name])
    # Run the C-program and capture the output
    # output = subprocess.run(, shell=True)
    subprocess.run([r".\a"], shell=True)
    time.sleep(0.5)
    img(ass_no,idx)
    os.system('cls')
    
res = {}
for i in range(2,8):
    res[i] = []
for i in range(2,8):
    
    path = f"G:\\Assignment-{i}" # Folder name where Assignment wise programs are allocated
    for p in os.listdir(path):
        if os.path.isfile(os.path.join(path, p)):
            if(p.endswith(".c")):
                if(p.startswith("p")):#Pls give the programs a numbered pattern like 1.c 2.c, etc...or write p1.c, p2.c otherwise it won't work
                    idx = int(p.lstrip("p").rstrip(".c"))
                    
                else:
                    idx = int(p.rstrip(".c"))
                run_c_program(i, idx,os.path.join(path, p))
                res[i].append([os.path.join(path, p), f"screenshot{i}({idx}).jpg", idx])
               
print(res)
input()               
doc_creator(res)


